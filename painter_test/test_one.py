import json

from docx.shared import Inches
from pyecharts import options as opts
from pyecharts.charts import Page, Radar, Line, Bar, Grid
import requests
from snapshot_selenium import snapshot as driver
from pyecharts.render import make_snapshot
from docx import Document

api1 = 'http://47.94.133.29/backstage_detail/warning_describe/?id=300266201708031&end_date=2018-09-03 '
api2 = 'http://47.94.133.29/backstage_detail/warning_describe/?id=300266201708031&end_date=2018-09-03 '


def get_api_dict(url):
    json_one = requests.get(url).text
    radar_dict_str = json.loads(json_one)['risk_radar']
    company_name = json.loads(json_one)["stock_name"]
    radar_dict = json.loads(radar_dict_str)
    return radar_dict, company_name


def radar_base(schema, value, company_name) -> Radar:
    c = (
        Radar()
            .add_schema(
            schema=schema,
            shape='circle',
            textstyle_opts=opts.TextStyleOpts(font_size=20),
        )
            .add(company_name, value, areastyle_opts=opts.AreaStyleOpts(color='red', opacity=0.5))
            .set_series_opts(label_opts=opts.LabelOpts(is_show=False))
            .set_global_opts(title_opts=opts.TitleOpts(title="雷达预警图"))
    )
    return c


name_dict = {
    "market_wave": "行情异动",
    "big_trans": "大宗交易",
    "holder_reduce": "股东减持",
    "equities_pledge": "股权质押",
    "market_news": "市场消息",
    "investor_wave": "投资者异动",
    "report_abnormal": "财报异常",
    "good_anno": "利好公告"
}

# 生成雷达图片
dict_one, company_name = get_api_dict(api1)
print('生成了数据1')
value_list = [[]]
c_schema = []
for key, value in dict_one.items():
    tmp = {}
    name = name_dict[key]
    tmp['name'] = name
    tmp['max'] = 1
    tmp['min'] = -1
    c_schema.append(tmp)
    value_list[0].append(value)

radar = radar_base(schema=c_schema, value=value_list, company_name=company_name)
make_snapshot(driver, radar.render(), 'radar.png')


# 获取折线图和柱状图
def grid_horizontal(name_list, value_list, company_name):
    line = (
        Line()
            .add_xaxis(name_list)
            .add_yaxis(company_name, value_list)
            .set_global_opts(
            title_opts=opts.TitleOpts(title="Grid-Line", pos_right="5%"),
            legend_opts=opts.LegendOpts(pos_right="20%"),
        )
    )
    bar = (
        Bar()
            .add_xaxis(name_list)
            .add_yaxis(company_name, value_list)
            .set_global_opts(title_opts=opts.TitleOpts(title="Bar-基本示例", subtitle="我是副标题"))
    )

    grid = (
        Grid()
            .add(line, grid_opts=opts.GridOpts(pos_left="55%"))
            .add(bar, grid_opts=opts.GridOpts(pos_right="55%"))
    )
    return grid


dict_two, company_name = get_api_dict(api2)
print('生成了数据2')
name_list2 = []
value_list2 = []

for key, value in dict_two.items():
    name_list2.append(name_dict[key])
    value_list2.append(value)

grid = grid_horizontal(name_list=name_list2, value_list=value_list2, company_name=company_name)
make_snapshot(driver, grid.render(), 'grid.png')

print("图片已经生成完毕")
# 将两张图片保存到word里

document = Document()
document.add_heading("数据图表展示")
document.add_picture('radar.png', width=Inches(5.0))
document.add_picture('grid.png', width=Inches(5.0))
document.save('one_file.docx')
print('文档生成完毕')
